"""
EMMA v2.0 — Module Documents & Rapports
Generation PDF (ReportLab), DOCX (python-docx), notes vocales.
Createur : ONANA GREGOIRE LEGRAND | Yaounde, Cameroun
"""

import asyncio
import logging
from datetime import datetime
from pathlib import Path

log = logging.getLogger("emma.actions.documents")

DOCS_DIR = Path.home() / "Documents" / "EMMA"


async def execute(sub_action: str, params: dict) -> str:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _dispatch, sub_action, params)


def _dispatch(sub_action: str, params: dict) -> str:
    handlers = {
        "report":      _report,
        "note":        _note,
        "summary":     _summary,
        "journal":     _journal,
        "prospection": _prospection,
    }
    fn = handlers.get(sub_action)
    return fn(params) if fn else f"Action document '{sub_action}' inconnue, Monsieur."


def _ensure_dir():
    DOCS_DIR.mkdir(parents=True, exist_ok=True)


# ─── Rapport d'activite ───────────────────────────────────────────────────────

def _report(params: dict) -> str:
    _ensure_dir()
    title    = params.get("title", "Rapport EMMA")
    content  = params.get("content", "Aucun contenu fourni.")
    fmt      = params.get("format", "docx").lower()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"rapport_{timestamp}.{fmt}"
    path      = DOCS_DIR / filename

    if fmt == "pdf":
        _write_pdf(str(path), title, content)
    else:
        _write_docx(str(path), title, content)

    return f"Rapport '{filename}' genere et sauvegarde, Monsieur."


def _write_docx(path: str, title: str, content: str):
    from docx import Document
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(path)
    log.info(f"[Documents] DOCX genere : {path}")


def _write_pdf(path: str, title: str, content: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    doc    = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    story  = [
        Paragraph(title, styles["Title"]),
        Paragraph(datetime.now().strftime('%d/%m/%Y %H:%M'), styles["Normal"]),
        Spacer(1, 12),
    ]
    for line in content.split("\n"):
        story.append(Paragraph(line, styles["Normal"]))
    doc.build(story)
    log.info(f"[Documents] PDF genere : {path}")


# ─── Note rapide ──────────────────────────────────────────────────────────────

def _note(params: dict) -> str:
    _ensure_dir()
    text      = params.get("text", "")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path      = DOCS_DIR / f"note_{timestamp}.txt"

    path.write_text(
        f"[{datetime.now().strftime('%d/%m/%Y %H:%M')}]\n{text}\n",
        encoding="utf-8"
    )
    return f"Note enregistree sous '{path.name}', Monsieur."


# ─── Synthese de recherche ────────────────────────────────────────────────────

def _summary(params: dict) -> str:
    content  = params.get("content", "")
    topic    = params.get("topic", "synthese")
    fmt      = params.get("format", "docx").lower()
    _ensure_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"synthese_{topic}_{timestamp}.{fmt}"
    path      = DOCS_DIR / filename

    if fmt == "pdf":
        _write_pdf(str(path), f"Synthese : {topic}", content)
    else:
        _write_docx(str(path), f"Synthese : {topic}", content)

    return f"Synthese '{filename}' generee, Monsieur."


# ─── Journal EMMA ─────────────────────────────────────────────────────────────

def _journal(params: dict) -> str:
    import json

    _ensure_dir()
    today     = datetime.now().strftime("%Y-%m-%d")
    log_path  = Path.home() / ".emma" / f"journal_{today}.json"

    if not log_path.exists():
        return "Aucun journal pour aujourd'hui, Monsieur."

    with open(log_path, "r", encoding="utf-8") as f:
        entries = json.load(f)

    out_path = DOCS_DIR / f"journal_{today}.txt"
    lines    = [f"JOURNAL EMMA - {today}\n{'='*40}\n"]
    for e in entries:
        lines.append(f"[{e['timestamp']}] {e['action']} — {e.get('result', '')}")
    out_path.write_text("\n".join(lines), encoding="utf-8")

    return f"Journal du {today} sauvegarde dans '{out_path.name}', Monsieur."


# ─── Dossier de prospection ───────────────────────────────────────────────────

def _prospection(params: dict) -> str:
    company = params.get("company", "Entreprise")
    content = params.get("content", "")
    _ensure_dir()
    path = DOCS_DIR / f"prospection_{company.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    _write_docx(str(path), f"Dossier de Prospection : {company}", content)
    return f"Dossier de prospection pour '{company}' genere, Monsieur."
